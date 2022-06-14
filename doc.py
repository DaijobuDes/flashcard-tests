import docx
from ppt import is_bold, is_term, is_underlined

# Documentation: https://python-docx.readthedocs.io/en/latest/

def main(data: docx.Document):
    text_runs = []

    # print(data.paragraphs)
    for i in data.paragraphs:
        for j in i.runs:
            if is_bold(j):
                text_runs.append(j.text)
            elif is_underlined(j):
                text_runs.append(j.text)

    return text_runs
            # print(j.text)
            # print(True if is_bold(j) else False)

    # for slide in data.slides:
    #     for shape in slide.shapes:
    #         if not shape.has_text_frame:
    #             continue

    #         for paragraph in shape.text_frame.paragraphs:
    #             for run in paragraph.runs:
    #                 if is_bold(run):
    #                     # print(f"Term: {run.text}")
    #                     text_runs.append(run.text)
    #                 elif is_underlined(run):
    #                     # print(f"Def: {run.text}")
    #                     text_runs.append(run.text)


if __name__ == '__main__':
    data = docx.Document(".\\test.docx")
    print(main(data))
